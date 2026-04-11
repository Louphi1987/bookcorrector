from __future__ import annotations

import io

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def build_docx_report(
    title: str,
    original_text: str,
    corrected_segments: list[dict[str, object]],
    issues: list[dict[str, object]],
    selected_ids: set[str],
) -> bytes:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    heading = document.add_heading("Karamaz Correcteur", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = document.add_paragraph()
    intro.add_run("Document source : ").bold = True
    intro.add_run(title)

    summary = document.add_paragraph()
    summary.add_run("Corrections appliquees : ").bold = True
    summary.add_run(str(len(selected_ids)))
    summary.add_run(" | Suggestions detectees : ").bold = True
    summary.add_run(str(len(issues)))

    document.add_heading("Version corrigee", level=1)
    _append_segmented_text(document, corrected_segments)

    document.add_section(WD_SECTION_START.NEW_PAGE)
    document.add_heading("Texte original", level=1)
    _append_original_text(document, original_text)

    document.add_heading("Rapport des suggestions", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Appliquee"
    headers[1].text = "Categorie"
    headers[2].text = "Extrait"
    headers[3].text = "Suggestion"
    headers[4].text = "Commentaire"

    for issue in issues:
        row = table.add_row().cells
        row[0].text = "Oui" if issue["id"] in selected_ids else "Non"
        row[1].text = str(issue["category"])
        row[2].text = str(issue["excerpt"])
        row[3].text = str(issue.get("replacement") or issue.get("suggestion") or "")
        row[4].text = str(issue.get("message") or "")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _append_segmented_text(document: Document, segments: list[dict[str, object]]) -> None:
    paragraph = document.add_paragraph()
    for segment in segments:
        text = str(segment["text"])
        if not text:
            continue
        chunks = text.split("\n")
        for index, chunk in enumerate(chunks):
            run = paragraph.add_run(chunk)
            if segment.get("changed"):
                run.bold = True
                run.font.color.rgb = RGBColor(0x9A, 0x34, 0x12)
            if index < len(chunks) - 1:
                paragraph = document.add_paragraph()


def _append_original_text(document: Document, original_text: str) -> None:
    for block in original_text.split("\n"):
        document.add_paragraph(block)
